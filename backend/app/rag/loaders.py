"""Load raw 3GPP specification files into normalised plain text.

3GPP publishes specs as .doc/.docx (and PDF renderings). We support PDF, DOCX
and plain text/markdown. The goal of this module is *only* to produce clean
text plus document-level metadata; clause structure is recovered later by
`chunking.py`, which is where the domain knowledge lives.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# "TS 23.501", "TR 38.913", "3GPP TS 24.501 V17.9.0"
_SPEC_ID_RE = re.compile(r"\b(TS|TR)\s?(\d{2}\.\d{3})\b", re.IGNORECASE)
_VERSION_RE = re.compile(r"\bV?(\d+\.\d+\.\d+)\b")
_RELEASE_RE = re.compile(r"\bRelease\s+(\d{1,2})\b", re.IGNORECASE)


@dataclass
class LoadedDocument:
    """A whole specification file, flattened to text."""

    source_path: str
    filename: str
    text: str
    doc_id: str               # e.g. "TS 23.501"
    title: str                # e.g. "System architecture for the 5G System"
    version: str = ""         # e.g. "17.9.0"
    release: str = ""         # e.g. "17"
    page_offsets: list = field(default_factory=list)  # char offset -> page no.

    def page_for_offset(self, offset: int) -> int | None:
        """Map a character offset back to a 1-indexed page (PDF sources only)."""
        if not self.page_offsets:
            return None
        page = 1
        for start, number in self.page_offsets:
            if offset >= start:
                page = number
            else:
                break
        return page


# ---------------------------------------------------------------------------
# Text normalisation
# ---------------------------------------------------------------------------

def _normalise(text: str) -> str:
    """Undo the usual PDF/Word extraction damage without destroying layout."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("‑", "-")
    text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")
    # De-hyphenate words broken across lines: "regis-\ntration" -> "registration"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # 3GPP uses tabs between clause number and title; keep a single space.
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse runs of blank lines but keep paragraph breaks meaningful.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_running_headers(text: str) -> str:
    """Remove per-page boilerplate that PDF extraction repeats hundreds of times.

    Any short line that appears on more than a third of the pages is treated as
    a running header/footer. Without this, retrieval is polluted by chunks that
    are 40% copyright notice.
    """
    lines = text.split("\n")
    counts: dict[str, int] = {}
    for line in lines:
        stripped = line.strip()
        if 0 < len(stripped) <= 90:
            counts[stripped] = counts.get(stripped, 0) + 1

    page_count = max(text.count("\f"), 1)
    threshold = max(4, page_count // 3)
    noisy = {
        line
        for line, n in counts.items()
        if n >= threshold
        and (
            re.match(r"^\d+$", line)                       # bare page numbers
            or "3GPP" in line
            or "ETSI" in line
            or re.match(r"^(TS|TR)\s?\d{2}\.\d{3}", line, re.I)
            or "Release" in line
        )
    }
    if not noisy:
        return text
    return "\n".join(l for l in lines if l.strip() not in noisy)


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------

def _extract_metadata(text: str, filename: str) -> tuple[str, str, str, str]:
    """Return (doc_id, title, version, release).

    Filename wins over content when both are available — 3GPP filenames are
    reliable (`23501-h90.docx`) while the cover page is often mangled by the
    extractor.
    """
    head = text[:4000]

    doc_id = ""
    m = _SPEC_ID_RE.search(filename) or _SPEC_ID_RE.search(head)
    if m:
        doc_id = f"{m.group(1).upper()} {m.group(2)}"
    else:
        # 3GPP archive naming: 23501-h90.docx -> TS 23.501
        m2 = re.match(r"^(\d{2})(\d{3})[-_.]", filename)
        if m2:
            doc_id = f"TS {m2.group(1)}.{m2.group(2)}"
    if not doc_id:
        doc_id = Path(filename).stem

    version = ""
    mv = _VERSION_RE.search(head)
    if mv:
        version = mv.group(1)

    release = ""
    mr = _RELEASE_RE.search(head)
    if mr:
        release = mr.group(1)
    elif version:
        release = version.split(".")[0]

    # Title: first substantial line that isn't the spec number or boilerplate.
    title = ""
    for line in head.split("\n"):
        s = line.strip()
        if len(s) < 12 or len(s) > 160:
            continue
        if _SPEC_ID_RE.search(s) or s.lower().startswith(("3gpp", "etsi", "release")):
            continue
        if re.match(r"^\d+(\.\d+)*\s", s):  # already a clause heading
            continue
        title = s
        break
    if not title:
        title = Path(filename).stem

    return doc_id, title, version, release


# ---------------------------------------------------------------------------
# Per-format readers
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> tuple[str, list]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pypdf is required to ingest PDF files") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    offsets: list[tuple[int, int]] = []
    cursor = 0
    for i, page in enumerate(reader.pages, start=1):
        content = page.extract_text() or ""
        offsets.append((cursor, i))
        parts.append(content)
        cursor += len(content) + 1
    return "\n".join(parts), offsets


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to ingest DOCX files") from exc

    document = docx.Document(str(path))
    blocks: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            blocks.append("")
            continue
        # Word heading styles carry the clause structure; re-emit the number so
        # the chunker's regex can see it even when Word auto-numbered it.
        style = (para.style.name or "").lower()
        if style.startswith("heading") and not re.match(r"^\d", text):
            blocks.append(text)
        else:
            blocks.append(text)

    # Tables matter a lot in 3GPP (QoS characteristics, timer values...).
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n" + "\n".join(rows) + "\n")

    return "\n".join(blocks)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_document(path: str | Path) -> LoadedDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type {suffix!r} for {path.name}")

    offsets: list = []
    if suffix == ".pdf":
        raw, offsets = _read_pdf(path)
    elif suffix == ".docx":
        raw = _read_docx(path)
    else:
        raw = _read_text(path)

    text = _normalise(_strip_running_headers(raw))
    doc_id, title, version, release = _extract_metadata(text, path.name)

    return LoadedDocument(
        source_path=str(path),
        filename=path.name,
        text=text,
        doc_id=doc_id,
        title=title,
        version=version,
        release=release,
        page_offsets=offsets,
    )


def discover_documents(corpus_dir: str | Path) -> list[Path]:
    corpus_dir = Path(corpus_dir)
    if not corpus_dir.exists():
        return []
    found = [
        p
        for p in sorted(corpus_dir.rglob("*"))
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES and not p.name.startswith(".")
    ]
    return found
